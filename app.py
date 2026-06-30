import os
import re
import streamlit as st

# 可选依赖：上传 Excel / Word 时会用到
try:
    import pandas as pd
except Exception:
    pd = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None


# =========================
# 基础路径
# =========================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
avatar_path = os.path.join(BASE_DIR, "avatar.png")
video_path = os.path.join(BASE_DIR, "demo_video.mp4")


# =========================
# 页面配置
# =========================
st.set_page_config(
    page_title="职联AI数字人带岗服务系统",
    page_icon="🤖",
    layout="wide"
)


# =========================
# 默认岗位信息
# =========================
DEFAULT_VALUES = {
    "job_name": "社区就业服务专员",
    "company": "某街道便民服务中心",
    "location": "北京市海淀区",
    "salary": "4000-6000元/月",
    "education": "大专及以上",
    "experience": "不限，有相关服务经验优先",
    "duty": "负责就业政策咨询、岗位信息登记、求职人员接待与基础材料整理。",
    "requirement": "沟通表达能力较好，责任心强，熟悉基础办公软件。",
    "welfare": "五险一金、节假日福利、工作稳定",
    "apply_method": "可前往便民服务中心咨询报名"
}

for key, value in DEFAULT_VALUES.items():
    if key not in st.session_state:
        st.session_state[key] = value


# =========================
# 工具函数：读取上传文档
# =========================
def read_txt_file(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    for encoding in ["utf-8", "gbk", "gb18030"]:
        try:
            return data.decode(encoding)
        except Exception:
            continue
    return data.decode("utf-8", errors="ignore")


def read_docx_file(uploaded_file) -> str:
    if Document is None:
        return "当前环境缺少 python-docx，无法解析 Word 文档。请在 requirements.txt 中加入 python-docx。"

    doc = Document(uploaded_file)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # 读取表格内容
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_texts.append(" | ".join(cells))

    return "\n".join(paragraphs + table_texts)


def read_excel_file(uploaded_file) -> str:
    if pd is None:
        return "当前环境缺少 pandas/openpyxl，无法解析 Excel 文档。请在 requirements.txt 中加入 pandas 和 openpyxl。"

    try:
        sheets = pd.read_excel(uploaded_file, sheet_name=None)
        all_text = []
        for sheet_name, df in sheets.items():
            all_text.append(f"【工作表：{sheet_name}】")
            all_text.append(df.astype(str).fillna("").to_string(index=False))
        return "\n".join(all_text)
    except Exception as e:
        return f"Excel 解析失败：{e}"


def read_pdf_file(uploaded_file) -> str:
    if PdfReader is None:
        return "当前环境缺少 pypdf，无法解析 PDF 文档。请在 requirements.txt 中加入 pypdf。"

    try:
        reader = PdfReader(uploaded_file)
        texts = []
        for page in reader.pages:
            text = page.extract_text() or ""
            texts.append(text)
        return "\n".join(texts)
    except Exception as e:
        return f"PDF 解析失败：{e}"


def read_uploaded_file(uploaded_file) -> str:
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".txt"):
        return read_txt_file(uploaded_file)

    if file_name.endswith(".docx"):
        return read_docx_file(uploaded_file)

    if file_name.endswith(".xlsx") or file_name.endswith(".xls"):
        return read_excel_file(uploaded_file)

    if file_name.endswith(".pdf"):
        return read_pdf_file(uploaded_file)

    return "暂不支持该文件格式。请上传 txt、docx、xlsx、xls 或 pdf 文件。"


# =========================
# 工具函数：从文本中提取岗位字段
# =========================
def find_by_patterns(text: str, patterns, default: str = "") -> str:
    for pattern in patterns:
        match = re.search(pattern, text, re.S)
        if match:
            value = match.group(1).strip()
            value = re.sub(r"\s+", " ", value)
            value = value.strip("：:;；,，。 ")
            if value and value.lower() not in ["nan", "none"]:
                return value
    return default


def extract_job_info(text: str) -> dict:
    clean_text = text.replace("\r", "\n")

    result = {
        "job_name": find_by_patterns(
            clean_text,
            [
                r"岗位名称[：:\s]+([^\n|]+)",
                r"招聘岗位[：:\s]+([^\n|]+)",
                r"职位名称[：:\s]+([^\n|]+)",
                r"岗位[：:\s]+([^\n|]+)"
            ],
            st.session_state["job_name"]
        ),
        "company": find_by_patterns(
            clean_text,
            [
                r"招聘单位[：:\s]+([^\n|]+)",
                r"用人单位[：:\s]+([^\n|]+)",
                r"单位名称[：:\s]+([^\n|]+)",
                r"公司名称[：:\s]+([^\n|]+)"
            ],
            st.session_state["company"]
        ),
        "location": find_by_patterns(
            clean_text,
            [
                r"工作地点[：:\s]+([^\n|]+)",
                r"工作地址[：:\s]+([^\n|]+)",
                r"办公地点[：:\s]+([^\n|]+)"
            ],
            st.session_state["location"]
        ),
        "salary": find_by_patterns(
            clean_text,
            [
                r"薪资待遇[：:\s]+([^\n|]+)",
                r"薪酬待遇[：:\s]+([^\n|]+)",
                r"工资待遇[：:\s]+([^\n|]+)",
                r"薪资[：:\s]+([^\n|]+)"
            ],
            st.session_state["salary"]
        ),
        "education": find_by_patterns(
            clean_text,
            [
                r"学历要求[：:\s]+([^\n|]+)",
                r"学历[：:\s]+([^\n|]+)"
            ],
            st.session_state["education"]
        ),
        "experience": find_by_patterns(
            clean_text,
            [
                r"工作经验[：:\s]+([^\n|]+)",
                r"经验要求[：:\s]+([^\n|]+)",
                r"从业经验[：:\s]+([^\n|]+)"
            ],
            st.session_state["experience"]
        ),
        "welfare": find_by_patterns(
            clean_text,
            [
                r"福利待遇[：:\s]+([^\n]+)",
                r"岗位福利[：:\s]+([^\n]+)",
                r"福利[：:\s]+([^\n]+)"
            ],
            st.session_state["welfare"]
        ),
        "apply_method": find_by_patterns(
            clean_text,
            [
                r"报名方式[：:\s]+([^\n]+)",
                r"联系方式[：:\s]+([^\n]+)",
                r"咨询方式[：:\s]+([^\n]+)"
            ],
            st.session_state["apply_method"]
        ),
    }

    duty = find_by_patterns(
        clean_text,
        [
            r"岗位职责[：:\s]+(.+?)(?:任职要求|岗位要求|资格要求|学历要求|薪资待遇|福利待遇|报名方式|$)",
            r"工作职责[：:\s]+(.+?)(?:任职要求|岗位要求|资格要求|学历要求|薪资待遇|福利待遇|报名方式|$)",
            r"工作内容[：:\s]+(.+?)(?:任职要求|岗位要求|资格要求|学历要求|薪资待遇|福利待遇|报名方式|$)"
        ],
        st.session_state["duty"]
    )

    requirement = find_by_patterns(
        clean_text,
        [
            r"任职要求[：:\s]+(.+?)(?:岗位职责|工作内容|薪资待遇|福利待遇|报名方式|联系方式|$)",
            r"岗位要求[：:\s]+(.+?)(?:岗位职责|工作内容|薪资待遇|福利待遇|报名方式|联系方式|$)",
            r"资格要求[：:\s]+(.+?)(?:岗位职责|工作内容|薪资待遇|福利待遇|报名方式|联系方式|$)"
        ],
        st.session_state["requirement"]
    )

    result["duty"] = duty
    result["requirement"] = requirement

    return result


# =========================
# 工具函数：生成内容
# =========================
def build_coze_input(job_info: dict) -> str:
    return f"""
请根据以下岗位信息，生成适合AI数字人直播/短视频口播的带岗内容。

岗位名称：{job_info["job_name"]}
招聘单位：{job_info["company"]}
工作地点：{job_info["location"]}
薪资待遇：{job_info["salary"]}
学历要求：{job_info["education"]}
工作经验：{job_info["experience"]}
岗位职责：{job_info["duty"]}
任职要求：{job_info["requirement"]}
福利待遇：{job_info["welfare"]}
报名方式：{job_info["apply_method"]}

输出要求：
1. 生成【直播开场白】；
2. 生成【岗位宣传文案】；
3. 生成【数字人口播稿】；
4. 生成【结尾引导语】；
5. 语言正式、亲和、清晰，符合政务就业服务表达风格；
6. 不夸大薪资待遇，不承诺录用结果；
7. 不使用“包过”“保证录用”“高薪急招”等表述；
8. 内容适合数字人直接朗读；
9. 最后加一句：内容由AI辅助生成，具体岗位信息以招聘单位官方发布为准。
""".strip()


def build_local_result(job_info: dict) -> dict:
    opening = f"""各位求职者朋友大家好，欢迎来到职联AI数字人带岗服务。本期为大家推荐的是「{job_info["job_name"]}」岗位。"""

    promo = f"""本次推荐岗位为「{job_info["job_name"]}」，招聘单位为{job_info["company"]}，工作地点位于{job_info["location"]}，薪资待遇为{job_info["salary"]}。该岗位学历要求为{job_info["education"]}，工作经验要求为{job_info["experience"]}。岗位主要面向有就业意向、愿意从事相关服务工作的求职者。"""

    script = f"""
各位求职者朋友大家好，欢迎来到职联AI数字人带岗直播间。

今天为大家推荐的岗位是：{job_info["job_name"]}。

本岗位招聘单位为{job_info["company"]}，工作地点在{job_info["location"]}，薪资待遇为{job_info["salary"]}。学历方面要求{job_info["education"]}，工作经验方面要求{job_info["experience"]}。

岗位主要职责包括：{job_info["duty"]}

任职要求方面，主要包括：{job_info["requirement"]}

福利待遇方面，岗位提供{job_info["welfare"]}。

有意向的求职者朋友，可以通过以下方式报名咨询：{job_info["apply_method"]}。

以上就是本期岗位推荐内容。建议大家结合自身学历、工作经验和求职意向进行选择，提前准备好个人简历及相关材料。感谢大家的关注，我们下期再见。

内容由AI辅助生成，具体岗位信息以招聘单位官方发布为准。
""".strip()

    ending = f"""有意向的求职者可通过「{job_info["apply_method"]}」进一步了解岗位详情，建议提前准备个人简历及相关材料。"""

    full_result = f"""
【直播开场白】
{opening}

【岗位宣传文案】
{promo}

【数字人口播稿】
{script}

【结尾引导语】
{ending}
""".strip()

    return {
        "opening": opening,
        "promo": promo,
        "script": script,
        "ending": ending,
        "full_result": full_result
    }


def build_video_prompt(job_info: dict, script_text: str) -> str:
    return f"""
请生成一条16:9横屏的政务就业服务类AI数字人带岗视频。

视频画面要求：
使用蓝衬衫女性AI数字人主播形象作为主要角色参考。画面为蓝色科技风直播间或政务服务直播间，数字人主播坐在桌前，面对镜头进行岗位介绍。整体风格正式、亲和、清晰，适合就业服务大厅展示和政务新媒体宣传。

角色要求：
女性AI数字人主播，蓝色衬衫，形象亲和，表情自然，坐姿端正，面对镜头播报。请尽量让人物口型与旁白内容同步，像直播带岗主播一样进行讲解。

岗位核心信息：
岗位名称：{job_info["job_name"]}
招聘单位：{job_info["company"]}
工作地点：{job_info["location"]}
薪资待遇：{job_info["salary"]}
学历要求：{job_info["education"]}
工作经验：{job_info["experience"]}
福利待遇：{job_info["welfare"]}

口播文案：
{script_text}

字幕要求：
自动添加中文字幕，字幕简洁清晰，重点信息可用信息卡形式展示，例如岗位名称、工作地点、薪资待遇、学历要求、福利待遇。

风格限制：
不要使用夸张营销风格，不要出现“包过”“保证录用”“高薪急招”等表述。内容要符合政务就业服务表达规范。
""".strip()


# =========================
# 页面标题
# =========================
st.title("职联AI数字人带岗服务系统")
st.caption("在线演示版｜岗位文档上传 → 岗位信息解析 → Coze话术生成 → 数字人口播视频展示")

st.markdown("""
本系统面向政务就业服务带岗场景，支持工作人员上传岗位信息文档或手动录入岗位信息，
系统可自动整理岗位要素，生成岗位宣传文案、直播开场白、数字人口播稿、视频生成提示词，
并展示AI数字人带岗视频样片。
""")


# =========================
# 一、岗位文档上传
# =========================
st.subheader("一、岗位文档上传与解析")

uploaded_file = st.file_uploader(
    "请上传岗位信息文档，支持 txt、docx、xlsx、xls、pdf",
    type=["txt", "docx", "xlsx", "xls", "pdf"]
)

if uploaded_file is not None:
    raw_text = read_uploaded_file(uploaded_file)

    with st.expander("查看文档解析原文", expanded=False):
        st.text_area("解析出的文档文本", raw_text, height=260)

    parsed_info = extract_job_info(raw_text)

    st.info("系统已尝试从文档中识别岗位信息。若识别结果不完整，可点击下方按钮填入表单后再手动微调。")

    with st.expander("查看识别结果", expanded=True):
        st.json(parsed_info, expanded=False)

    if st.button("使用识别结果填入岗位表单"):
        for key, value in parsed_info.items():
            if value:
                st.session_state[key] = value
        st.success("已将识别结果填入表单。请检查并微调后继续生成内容。")


# =========================
# 二、岗位信息录入 + 数字人形象
# =========================
left, right = st.columns([1.25, 0.75])

with left:
    st.subheader("二、岗位信息录入")

    st.text_input("岗位名称", key="job_name")
    st.text_input("招聘单位", key="company")
    st.text_input("工作地点", key="location")
    st.text_input("薪资待遇", key="salary")
    st.text_input("学历要求", key="education")
    st.text_input("工作经验", key="experience")
    st.text_area("岗位职责", key="duty", height=110)
    st.text_area("任职要求", key="requirement", height=110)
    st.text_input("福利待遇", key="welfare")
    st.text_input("报名方式", key="apply_method")

with right:
    st.subheader("三、AI数字人形象")

    try:
        with open(avatar_path, "rb") as img_file:
            avatar_bytes = img_file.read()
        st.image(avatar_bytes, caption="AI数字人带岗主播", use_column_width=True)
    except Exception as e:
        st.warning("未检测到 avatar.png，线上部署时请将 avatar.png 与 app.py 放在同一目录。")
        st.caption(f"图片加载信息：{e}")

    st.markdown("""
    **形象设定：**  
    女性数字人主播，蓝色衬衫，政务服务风格，语气亲和，表达规范。
    """)


# 当前岗位信息
job_info = {
    "job_name": st.session_state["job_name"],
    "company": st.session_state["company"],
    "location": st.session_state["location"],
    "salary": st.session_state["salary"],
    "education": st.session_state["education"],
    "experience": st.session_state["experience"],
    "duty": st.session_state["duty"],
    "requirement": st.session_state["requirement"],
    "welfare": st.session_state["welfare"],
    "apply_method": st.session_state["apply_method"],
}

local_result = build_local_result(job_info)
coze_input = build_coze_input(job_info)
video_prompt = build_video_prompt(job_info, local_result["script"])


# =========================
# 四、内容生成
# =========================
st.divider()
st.subheader("四、内容生成")

col1, col2, col3 = st.columns([1, 1, 1])

with col1:
    generate_btn = st.button("一键生成带岗内容", type="primary")

with col2:
    st.download_button(
        label="下载本地生成结果TXT",
        data=local_result["full_result"],
        file_name=f"{job_info['job_name']}_数字人带岗口播稿.txt",
        mime="text/plain"
    )

with col3:
    st.download_button(
        label="下载视频生成提示词TXT",
        data=video_prompt,
        file_name=f"{job_info['job_name']}_视频生成提示词.txt",
        mime="text/plain"
    )

if generate_btn:
    st.success("生成完成。当前版本已生成本地带岗内容、Coze输入内容和视频生成提示词。")

    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "直播开场白",
        "岗位宣传文案",
        "数字人口播稿",
        "结尾引导语",
        "Coze输入内容",
        "视频生成提示词"
    ])

    with tab1:
        st.write(local_result["opening"])

    with tab2:
        st.write(local_result["promo"])

    with tab3:
        st.text_area("可复制到数字人工具中的口播稿", local_result["script"], height=360)

    with tab4:
        st.write(local_result["ending"])

    with tab5:
        st.text_area("复制以下内容到 Coze / 扣子进行话术生成", coze_input, height=420)

    with tab6:
        st.text_area("复制以下内容到扣子视频 / 数字人视频工具", video_prompt, height=520)


# =========================
# 五、Coze 生成结果回填
# =========================
st.divider()
st.subheader("五、Coze 生成结果回填")

st.markdown("""
将 Coze / 扣子生成的带岗内容复制到下方文本框中，便于在网页端统一展示、保存和下载。
如果后续接入 Coze API，这一步可以改成系统自动调用并回填。
""")

coze_result = st.text_area(
    "请粘贴 Coze 生成的完整带岗内容",
    height=360,
    placeholder="在这里粘贴 Coze 生成的【直播开场白】【岗位宣传文案】【数字人口播稿】【结尾引导语】..."
)

if coze_result:
    st.success("已接收 Coze 生成结果，可用于数字人口播或短视频制作。")

    st.download_button(
        label="下载 Coze 生成结果TXT",
        data=coze_result,
        file_name=f"{job_info['job_name']}_Coze生成带岗内容.txt",
        mime="text/plain"
    )


# =========================
# 六、数字人视频样片
# =========================
st.divider()
st.subheader("六、数字人带岗视频样片")

st.markdown("""
下方展示由扣子视频生成的数字人口播测试片段，用于验证“AI数字人形象 + 口播内容 + 口型同步”的视频生成效果。
当前样片为短片段演示版，后续可接入视频生成接口，实现上传文档后一键生成完整数字人带岗视频。
""")

try:
    with open(video_path, "rb") as video_file:
        video_bytes = video_file.read()
    st.video(video_bytes)
    st.success("已加载数字人带岗演示视频。")
except Exception:
    st.info("如需展示视频样片，请将扣子生成的视频命名为 demo_video.mp4，并与 app.py 放在同一目录。")


# =========================
# 七、系统工作流程
# =========================
st.divider()
st.subheader("七、系统工作流程")

st.markdown("""
**第一步：岗位文档上传**  
工作人员上传 Word、Excel、PDF 或 TXT 格式的岗位信息文档，系统自动解析岗位名称、招聘单位、工作地点、薪资待遇、学历要求、岗位职责等基础字段。

**第二步：岗位信息确认**  
系统将识别出的岗位信息填入表单，工作人员可根据实际情况进行核对和微调。

**第三步：AI 带岗内容生成**  
系统根据岗位信息生成直播开场白、岗位宣传文案、数字人口播稿、结尾引导语，并整理出 Coze / 扣子输入内容。

**第四步：数字人视频生成验证**  
系统生成适用于扣子视频或数字人平台的视频提示词。当前版本已完成短片段数字人口播效果验证，后续可通过视频生成接口实现自动生成。

**第五步：带岗宣传应用**  
生成的内容可用于就业服务大厅展示、政务新媒体短视频、线上直播带岗和岗位推荐宣传。
""")


# =========================
# 八、后续升级方向
# =========================
st.divider()
st.subheader("八、后续升级方向")

st.markdown("""
1. **接入 Coze API**：实现网页端一键调用 Coze 工作流，自动返回带岗话术。  
2. **接入数字人视频 API**：实现上传岗位文档后一键生成口型同步数字人视频。  
3. **支持批量岗位生成**：上传 Excel 岗位表后批量生成多岗位带岗稿。  
4. **增加审核机制**：生成内容先由人工审核确认，再进入视频生成环节。  
5. **增加历史记录**：保存每次岗位文档、生成稿件和视频结果，方便后续查询复用。
""")